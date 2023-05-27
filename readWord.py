import os
from docx import Document
import re
import time
import pandas as pd
from openpyxl import load_workbook

FOLDER_PATH = "C:/Users/lxj/Documents/文理作业/asdf/"   # 请将此路径替换为Word的文件夹路径
SAVE_PATH = "C:/Users/lxj/Documents/文理作业/asdf/"    # 请将此路径替换为Excel要保存的文件夹路径


def countKeyword(docxFileName, filePath):   # 匹配关键字方法
    global isCriminalFlag
    keywords = {'出生': "出生日期", '族': "民族", '文化': "文化水平"}  # 字典关键字:key做endwith匹配, value做保存到excel的关键字
    document = Document(filePath)

    for para in document.paragraphs:  # 文档段落处理

        firstParagraphPattern = r'(上诉人|被告人)[^。]*'  # 首段匹配规则
        firstParagraphMatch = re.search(firstParagraphPattern, para.text)  # 一般来说第一段陈述了被告的基本信息，先筛选出第一段，然后再提取
        if firstParagraphMatch:
            keywordPattern = r'([^，]+)(族|出生|文化)'  # 正则匹配 民族、出生日期、文化水平
            mathRet = re.findall(keywordPattern, firstParagraphMatch.group())
            temp = ""
            for matchTuple in mathRet:
                for matchStr in matchTuple:
                    temp += matchStr  # 拼接元组字符串
                for keyword in keywords:
                    if temp.strip().endswith(keyword):  # 关键字
                        # print(f'keyword :{keywords.get(keyword)} string:{temp}')
                        data.append((docxFileName, keywords.get(keyword), temp))  # 添加匹配到的 民族、出生日期、文化水平 数据
                temp = ""

            # 单独匹配性别
            genderPattern = r'[男|女]'
            matchGender = re.findall(genderPattern, firstParagraphMatch.group())
            if matchGender:
                data.append((docxFileName, 'gender', *matchGender))

        # 法院判决段落
        verdictParagraphPattern = r'.*本院认为.*'
        verdictMatch = re.findall(verdictParagraphPattern, para.text)
        amountPattern = r'(数额[^.*]*?)，'  # 金额匹配规则
        criminalPattern = r'(曾因犯[^.*]*?罪[^.*]*?)，'  # 是否累犯匹配规则

        for match in verdictMatch:
            criminalMatches = re.search(criminalPattern, match)
            if criminalMatches:
                data.append((docxFileName, '是否累犯', criminalMatches.group(1)))
                isCriminalFlag = True
            stealAmountMatch = re.search(amountPattern, match)  # 盗窃数额性质
            if stealAmountMatch:
                data.append((docxFileName, '盗窃数额', stealAmountMatch.group(1)))
            # print(criminalMatches.group(1))

    # Document库 分割了所有段落，因此合并整篇文档用于匹配刑期
    docx = '\n'.join([para.text for para in document.paragraphs])  # 将整个文档的内容合并为一个字符串
    # 刑期判断
    verdictParagraphPattern2 = r'.*本院认为[^.*]*?\W*判决如下([\s\S]*?)审判'
    verdictMatch2 = re.search(verdictParagraphPattern2, docx)
    # print(para.text)
    if verdictMatch2:

        prisonTermPattern = r'犯[^。]*判[^。]*。'
        # print(verdictMatch2.group(1))
        prisonTermMatch = re.search(prisonTermPattern, verdictMatch2.group(1))
        if prisonTermMatch:
            # print(f'刑期：{prisonTermMatch.group()}')
            data.append((docxFileName, '刑期', prisonTermMatch.group()))
    docx = ''


def writeExcel(columns, rows, fileName):  # 判断是否累犯
    fileName = SAVE_PATH + fileName + ".xlsx"  # 路径

    if isCriminalFlag:
        appendWriter(columns, rows, fileName, 0)
    else:
        appendWriter(columns, rows, fileName, 1)


def appendWriter(columns, rows, fileName, sheetNum):  # 追加写入新数据方法
    # 加载 Excel 文件
    book = load_workbook(fileName)

    # 获取工作表
    writer = pd.ExcelWriter(fileName, engine='openpyxl')
    writer.book = book
    if sheetNum >= len(book.worksheets):
        # 创建新的工作表
        book.create_sheet(f'Sheet{(sheetNum.numerator + 1)}')
        book.worksheets[sheetNum].append(columns)
        writer.save()

    # 获取指定的工作表
    # print(f'当前工作表数量:length {len(book.worksheets)}')
    # print(f'当前写入的工作表:sheetNum {sheetNum}')
    sheet = book.worksheets[sheetNum]

    for row in rows:
        # print(row)
        sheet.append(row)
    # 保存文件
    writer.save()


def createFile(columns, fileName):  # 创建excel文件
    fileName = SAVE_PATH + fileName + ".xlsx"  # 路径
    # 将行数据转换为 DataFrame 对象, 实际上就是创建了一个excel文件,主要的数据写入在appendWriter
    # print(f"创建文件{fileName}")
    df = pd.DataFrame([], columns=columns)
    with pd.ExcelWriter(fileName) as writer:
        df.to_excel(writer, index=False)


if __name__ == '__main__':
    # 所有需检索的word文件路径
    # fileList = [
    #         'C:/Users/lxj/Documents/文理作业/大三下/网络安全/第六次实验/网安第六次.docx',
    #         'C:/Users/lxj/Documents/文理作业/asdf/Test.docx',
    # ]
    # for string in fileList:
    #     docxName = re.search(docxNamePattern, string).group(1)
    #     # data = countKeyword(data, docxName, string)
    #     countKeyword(data, docxName, string)
    #     # print(data)
    #     # writeExcel(header, data, csvFileName)
    #     data.clear()

    folder_path = FOLDER_PATH

    docxNamePattern = r'([^\/]+\.(docx|doc))(?!.*\/)'  # word文档名
    csvFileName = "" + time.strftime("%Y-%m-%d %H-%M-%S", time.localtime(time.time()))  # 以当前时间戳为命名
    header = ("fileName", "keyword", "content")  # 表头
    # header = ("fileName", *[keyword for keyword in keywords], "count")  # 表头
    data = []
    isCriminalFlag = False
    createFile(header, csvFileName)  # 调用创建文件方法.

    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)
        # print(filepath)
        if filepath.endswith(".docx") or filepath.endswith(".doc"):
            nameMatch = re.search(docxNamePattern, filepath)
            if nameMatch:
                docxName = nameMatch.group(1)
                # print(docxName)
                countKeyword(docxName, filepath)
                print(data)
                writeExcel(header, data, csvFileName)

                data.clear()
                # print("data clear ")

    print("done")
